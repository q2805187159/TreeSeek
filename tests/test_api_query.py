from __future__ import annotations

from pathlib import Path

import fitz
from docx import Document
from fastapi.testclient import TestClient

from app.main import app

client = TestClient(app)


def _create_pdf(path: Path, text: str):
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), text, fontsize=12)
    doc.save(path)
    doc.close()


def test_health():
    response = client.get("/health")
    assert response.status_code == 200
    assert response.json()["status"] == "ok"


def test_build_index_and_query_markdown(local_tmp_path):
    md_path = local_tmp_path / "sample.md"
    md_path.write_text("# Guide\n\n## Retrieval Design\n\nRetrieval design guidance.\n", encoding="utf-8")

    build_response = client.post(
        "/build-index",
        json={
            "path": str(md_path),
            "doc_type": "markdown",
            "include_text": True,
            "index_output_dir": str(local_tmp_path / "out"),
        },
    )
    assert build_response.status_code == 200
    payload = build_response.json()
    assert payload["index_path"]

    query_response = client.post(
        "/query",
        json={
            "index_path": payload["index_path"],
            "query": "retrieval design",
            "top_k": 5,
            "leaf_only": False,
            "debug_explain": True,
            "rerank_with_llm": False,
        },
    )
    assert query_response.status_code == 200
    assert query_response.json()["results"]


def test_build_corpus_and_query_corpus(local_tmp_path):
    docs_dir = local_tmp_path / "docs"
    docs_dir.mkdir()
    (docs_dir / "guide.md").write_text("# Guide\n\nRetrieval design guidance.\n", encoding="utf-8")
    _create_pdf(docs_dir / "report.pdf", "Financial stability and retrieval design.")
    word = Document()
    word.add_heading("Operations Manual", level=1)
    word.add_paragraph("Control testing and retrieval design.")
    word.save(docs_dir / "manual.docx")

    build_response = client.post(
        "/build-corpus",
        json={"input_dir": str(docs_dir), "corpus_name": "demo"},
    )
    assert build_response.status_code == 200
    payload = build_response.json()
    assert payload["document_count"] == 3

    query_response = client.post(
        "/query-corpus",
        json={
            "corpus_index_path": payload["corpus_index_path"],
            "query": "retrieval design",
            "top_k": 5,
            "doc_type": None,
            "tags": [],
            "source": None,
            "leaf_only": False,
            "debug_explain": False,
            "rerank_with_llm": False,
        },
    )
    assert query_response.status_code == 200
    assert query_response.json()["results"]


def test_build_corpus_respects_exclude_globs(local_tmp_path):
    docs_dir = local_tmp_path / "docs"
    docs_dir.mkdir()
    (docs_dir / "guide.md").write_text("# Guide\n\nRetrieval design guidance.\n", encoding="utf-8")
    (docs_dir / "MANUAL_TEST_STEPS.md").write_text("# Ignore\n\nDo not index me.\n", encoding="utf-8")

    response = client.post(
        "/build-corpus",
        json={
            "input_dir": str(docs_dir),
            "corpus_name": "demo-exclude",
            "exclude_globs": ["MANUAL_TEST_STEPS.md"],
        },
    )
    assert response.status_code == 200
    assert response.json()["document_count"] == 1
