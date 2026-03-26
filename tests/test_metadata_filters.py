from pathlib import Path

import fitz
from docx import Document

from treeseek import build_corpus_from_directory, search_corpus
from treeseek.corpus.corpus_models import CorpusQueryRequest


def _create_pdf(path: Path, text: str):
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), text, fontsize=12)
    doc.save(path)
    doc.close()


def _patch_pdf_builder(monkeypatch):
    def fake_build_pdf_tree_from_opt(path, opt=None):
        doc = fitz.open(path)
        texts = [page.get_text() for page in doc]
        doc.close()
        return {
            "doc_name": Path(path).name,
            "structure": [
                {
                    "title": Path(path).stem,
                    "node_id": "0000",
                    "start_index": 1,
                    "end_index": len(texts),
                    "text": "\n".join(texts),
                    "summary": texts[0][:200] if texts else "",
                }
            ],
        }

    monkeypatch.setattr("treeseek.corpus.corpus_builder.build_pdf_tree_from_opt", fake_build_pdf_tree_from_opt)


def test_metadata_filters_doc_type_and_source(local_tmp_path, monkeypatch):
    _patch_pdf_builder(monkeypatch)
    docs_dir = local_tmp_path / "docs"
    docs_dir.mkdir()
    (docs_dir / "guide.md").write_text("# Guide\n\nRetrieval design guidance.\n", encoding="utf-8")
    _create_pdf(docs_dir / "report.pdf", "Retrieval design in annual report.")
    word = Document()
    word.add_heading("Operations Manual", level=1)
    word.add_paragraph("Retrieval design for runbooks.")
    word.save(docs_dir / "manual.docx")

    corpus_index, _ = build_corpus_from_directory(
        str(docs_dir),
        corpus_name="demo",
        output_dir=str(local_tmp_path / "out"),
        user_opt={"if_add_node_summary": "no", "if_add_node_text": "yes", "index_include_text": "yes"},
    )

    pdf_results = search_corpus(
        corpus_index,
        CorpusQueryRequest(query="retrieval design", top_k=5, doc_type="pdf"),
    )
    assert pdf_results
    assert all(result.doc_type == "pdf" for result in pdf_results)

    source_results = search_corpus(
        corpus_index,
        CorpusQueryRequest(query="retrieval design", top_k=5, source="docs"),
    )
    assert source_results
    assert all(result.source == "docs" for result in source_results)
