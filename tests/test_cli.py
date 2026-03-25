from __future__ import annotations

import json
import os
import subprocess
import sys
from pathlib import Path


def test_markdown_cli_builds_query_index_and_runs_query(local_tmp_path):
    repo_root = Path(__file__).resolve().parents[1]
    md_path = local_tmp_path / "sample.md"
    index_dir = local_tmp_path / "index"
    results_dir = local_tmp_path / "results"

    md_path.write_text(
        "# Root\n\n"
        "Overview text.\n\n"
        "## Direct-to-Consumer\n\n"
        "Direct-to-consumer subscriber growth improved in Q1.\n\n"
        "## Parks\n\n"
        "Theme park attendance remained resilient.\n",
        encoding="utf-8",
    )

    command = [
        sys.executable,
        str(repo_root / "run_treeseek.py"),
        "--md_path",
        str(md_path),
        "--if-add-node-summary",
        "no",
        "--build-query-index",
        "yes",
        "--query",
        "direct-to-consumer",
        "--top-k",
        "2",
        "--include-text",
        "yes",
        "--index-output-dir",
        str(index_dir),
    ]

    completed = subprocess.run(
        command,
        cwd=local_tmp_path,
        capture_output=True,
        text=True,
        env={**os.environ, "PYTEST_DISABLE_PLUGIN_AUTOLOAD": "1"},
        check=True,
    )

    assert (results_dir / "sample_structure.json").exists()
    assert (index_dir / "sample_query_index.pkl.gz").exists()
    assert '"query": "direct-to-consumer"' in completed.stdout
    assert "Direct-to-Consumer" in completed.stdout
    assert '"snippet"' in completed.stdout
    assert '"snippet_field"' in completed.stdout
    assert '"field_scores"' not in completed.stdout


def test_query_only_mode_uses_existing_index_without_source_document(local_tmp_path):
    repo_root = Path(__file__).resolve().parents[1]
    md_path = local_tmp_path / "sample.md"
    index_dir = local_tmp_path / "index"
    index_path = index_dir / "sample_query_index.pkl.gz"

    md_path.write_text(
        "# Root\n\n"
        "Overview text.\n\n"
        "## Direct-to-Consumer\n\n"
        "Direct-to-consumer subscriber growth improved in Q1.\n\n",
        encoding="utf-8",
    )

    build_command = [
        sys.executable,
        str(repo_root / "run_treeseek.py"),
        "--md_path",
        str(md_path),
        "--if-add-node-summary",
        "no",
        "--build-query-index",
        "yes",
        "--include-text",
        "yes",
        "--index-output-dir",
        str(index_dir),
    ]
    subprocess.run(
        build_command,
        cwd=local_tmp_path,
        capture_output=True,
        text=True,
        env={**os.environ, "PYTEST_DISABLE_PLUGIN_AUTOLOAD": "1"},
        check=True,
    )

    md_path.unlink()

    query_command = [
        sys.executable,
        str(repo_root / "run_treeseek.py"),
        "--index-path",
        str(index_path),
        "--query",
        "direct-to-consumer",
        "--top-k",
        "2",
    ]
    completed = subprocess.run(
        query_command,
        cwd=local_tmp_path,
        capture_output=True,
        text=True,
        env={**os.environ, "PYTEST_DISABLE_PLUGIN_AUTOLOAD": "1"},
        check=True,
    )

    assert index_path.exists()
    assert '"query": "direct-to-consumer"' in completed.stdout
    assert "Direct-to-Consumer" in completed.stdout
    assert '"snippet"' in completed.stdout
    assert '"highlight_terms"' in completed.stdout
    assert '"field_scores"' not in completed.stdout


def test_debug_explain_mode_includes_explain_fields(local_tmp_path):
    repo_root = Path(__file__).resolve().parents[1]
    md_path = local_tmp_path / "sample.md"
    index_dir = local_tmp_path / "index"
    index_path = index_dir / "sample_query_index.pkl.gz"

    md_path.write_text(
        "# Root\n\n"
        "## Retrieval Design\n\n"
        "Retrieval design uses deterministic candidate generation.\n\n"
        "## Logging\n\n"
        "Logging and observability improve diagnostics.\n",
        encoding="utf-8",
    )

    subprocess.run(
        [
            sys.executable,
            str(repo_root / "run_treeseek.py"),
            "--md_path",
            str(md_path),
            "--if-add-node-summary",
            "no",
            "--build-query-index",
            "yes",
            "--include-text",
            "yes",
            "--index-output-dir",
            str(index_dir),
        ],
        cwd=local_tmp_path,
        capture_output=True,
        text=True,
        env={**os.environ, "PYTEST_DISABLE_PLUGIN_AUTOLOAD": "1"},
        check=True,
    )

    completed = subprocess.run(
        [
            sys.executable,
            str(repo_root / "run_treeseek.py"),
            "--index-path",
            str(index_path),
            "--query",
            "\"retrieval design\"",
            "--top-k",
            "2",
            "--debug-explain",
            "yes",
        ],
        cwd=local_tmp_path,
        capture_output=True,
        text=True,
        env={**os.environ, "PYTEST_DISABLE_PLUGIN_AUTOLOAD": "1"},
        check=True,
    )

    assert '"field_scores"' in completed.stdout
    assert '"bonuses_applied"' in completed.stdout
    assert '"phrase_matches"' in completed.stdout


def test_docx_cli_builds_query_index_and_runs_query(local_tmp_path):
    from docx import Document

    repo_root = Path(__file__).resolve().parents[1]
    docx_path = local_tmp_path / "sample.docx"
    index_dir = local_tmp_path / "index"
    results_dir = local_tmp_path / "results"

    document = Document()
    document.add_heading("Operations Guide", level=1)
    document.add_paragraph("Operational overview.")
    document.add_heading("Compliance Workflow", level=2)
    document.add_paragraph("Compliance review workflow and control testing.")
    document.save(docx_path)

    command = [
        sys.executable,
        str(repo_root / "run_treeseek.py"),
        "--docx_path",
        str(docx_path),
        "--if-add-node-summary",
        "no",
        "--build-query-index",
        "yes",
        "--query",
        "compliance workflow",
        "--top-k",
        "2",
        "--include-text",
        "yes",
        "--index-output-dir",
        str(index_dir),
    ]

    completed = subprocess.run(
        command,
        cwd=local_tmp_path,
        capture_output=True,
        text=True,
        env={**os.environ, "PYTEST_DISABLE_PLUGIN_AUTOLOAD": "1"},
        check=True,
    )

    assert (results_dir / "sample_structure.json").exists()
    assert (index_dir / "sample_query_index.pkl.gz").exists()
    assert '"query": "compliance workflow"' in completed.stdout
    assert "Compliance Workflow" in completed.stdout
