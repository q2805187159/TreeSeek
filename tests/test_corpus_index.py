import json
from pathlib import Path

import fitz
from docx import Document

from treeseek import build_corpus_from_directory, search_corpus
from treeseek.corpus.corpus_models import CorpusQueryRequest


def _create_pdf(path: Path, text: str):
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), text, fontsize=12)
    doc.save(path)
    doc.close()


def _patch_pdf_builder(monkeypatch):
    def fake_build_pdf_tree_from_opt(path, opt=None):
        doc = fitz.open(path)
        texts = [page.get_text() for page in doc]
        doc.close()
        return {
            "doc_name": Path(path).name,
            "structure": [
                {
                    "title": Path(path).stem,
                    "node_id": "0000",
                    "start_index": 1,
                    "end_index": len(texts),
                    "text": "\n".join(texts),
                    "summary": texts[0][:200] if texts else "",
                }
            ],
        }

    monkeypatch.setattr("treeseek.corpus.corpus_builder.build_pdf_tree_from_opt", fake_build_pdf_tree_from_opt)


def test_build_corpus_from_directory_creates_index_and_records(local_tmp_path, monkeypatch):
    _patch_pdf_builder(monkeypatch)
    docs_dir = local_tmp_path / "docs"
    docs_dir.mkdir()

    (docs_dir / "guide.md").write_text("# Guide\n\n## Retrieval Design\n\nRetrieval design guidance.\n", encoding="utf-8")
    _create_pdf(docs_dir / "report.pdf", "Financial stability and supervision report.")
    word = Document()
    word.add_heading("Operations Manual", level=1)
    word.add_paragraph("Control testing and compliance workflow.")
    word.save(docs_dir / "manual.docx")

    corpus_index, corpus_index_path = build_corpus_from_directory(
        str(docs_dir),
        corpus_name="demo",
        output_dir=str(local_tmp_path / "out"),
        user_opt={"if_add_node_summary": "no", "if_add_node_text": "yes", "index_include_text": "yes"},
    )

    assert len(corpus_index.documents) == 3
    assert Path(corpus_index_path).exists()
    for record in corpus_index.documents:
        assert Path(record.query_index_path).exists()
        assert Path(local_tmp_path / "out" / "demo" / f"{record.doc_id}_structure.json").exists()


def test_search_corpus_returns_cross_document_results(local_tmp_path, monkeypatch):
    _patch_pdf_builder(monkeypatch)
    docs_dir = local_tmp_path / "docs"
    docs_dir.mkdir()
    (docs_dir / "guide.md").write_text("# Guide\n\n## Retrieval Design\n\nRetrieval design guidance.\n", encoding="utf-8")
    _create_pdf(docs_dir / "report.pdf", "Financial stability and retrieval design.")

    corpus_index, _ = build_corpus_from_directory(
        str(docs_dir),
        corpus_name="demo",
        output_dir=str(local_tmp_path / "out"),
        user_opt={"if_add_node_summary": "no", "if_add_node_text": "yes", "index_include_text": "yes"},
    )

    results = search_corpus(
        corpus_index,
        CorpusQueryRequest(query="retrieval design", top_k=5, debug_explain=True),
    )

    assert results
    assert all(result.doc_id for result in results)
    assert all(result.doc_name for result in results)


def test_build_corpus_from_directory_respects_exclude_globs(local_tmp_path, monkeypatch):
    _patch_pdf_builder(monkeypatch)
    docs_dir = local_tmp_path / "docs"
    docs_dir.mkdir()

    (docs_dir / "guide.md").write_text("# Guide\n\nRetrieval design guidance.\n", encoding="utf-8")
    (docs_dir / "MANUAL_TEST_STEPS.md").write_text("# Ignore\n\nDo not index me.\n", encoding="utf-8")
    _create_pdf(docs_dir / "report.pdf", "Financial stability and supervision report.")

    corpus_index, _ = build_corpus_from_directory(
        str(docs_dir),
        corpus_name="demo",
        output_dir=str(local_tmp_path / "out"),
        user_opt={"if_add_node_summary": "no", "if_add_node_text": "yes", "index_include_text": "yes"},
        exclude_globs=["MANUAL_TEST_STEPS.md"],
    )

    assert len(corpus_index.documents) == 2
    assert "MANUAL_TEST_STEPS.md" not in {record.doc_name for record in corpus_index.documents}
