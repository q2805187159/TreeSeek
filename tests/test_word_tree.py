import asyncio

from treeseek.indexing.builder import build_query_index
from treeseek.indexing.query_engine import search_index
from treeseek.word_tree import build_word_tree


def test_build_word_tree_from_headings(local_tmp_path):
    from docx import Document

    docx_path = local_tmp_path / "sample.docx"
    document = Document()
    document.add_heading("Root Section", level=1)
    document.add_paragraph("Overview text for the root section.")
    document.add_heading("Child Section", level=2)
    document.add_paragraph("Detailed retrieval workflow and diagnostics.")
    document.save(docx_path)

    result = asyncio.run(
        build_word_tree(
            docx_path=str(docx_path),
            if_add_node_summary="no",
            if_add_node_text="yes",
            if_add_node_id="yes",
        )
    )

    assert result["doc_name"] == "sample.docx"
    assert result["structure"]
    assert result["structure"][0]["title"] == "Root Section"
    assert result["structure"][0]["nodes"][0]["title"] == "Child Section"
    assert "Detailed retrieval workflow" in result["structure"][0]["nodes"][0]["text"]


def test_build_word_tree_without_headings_falls_back_to_single_root(local_tmp_path):
    from docx import Document

    docx_path = local_tmp_path / "plain.docx"
    document = Document()
    document.add_paragraph("This document has no heading styles.")
    document.add_paragraph("It should fall back to a single root node.")
    document.save(docx_path)

    result = asyncio.run(
        build_word_tree(
            docx_path=str(docx_path),
            if_add_node_summary="no",
            if_add_node_text="yes",
            if_add_node_id="yes",
        )
    )

    assert len(result["structure"]) == 1
    assert result["structure"][0]["title"] == "plain"
    assert "single root node" in result["structure"][0]["text"]


def test_word_title_style_becomes_non_leaf_root(local_tmp_path):
    from docx import Document

    docx_path = local_tmp_path / "title_root.docx"
    document = Document()
    document.add_heading("Document Title", level=0)
    document.add_heading("Main Section", level=1)
    document.add_paragraph("Main section content.")
    document.add_heading("Child Section", level=2)
    document.add_paragraph("Control testing and compliance workflow.")
    document.save(docx_path)

    result = asyncio.run(
        build_word_tree(
            docx_path=str(docx_path),
            if_add_node_summary="no",
            if_add_node_text="yes",
            if_add_node_id="yes",
        )
    )

    assert result["structure"][0]["title"] == "Document Title"
    assert result["structure"][0]["nodes"][0]["title"] == "Main Section"
    index = build_query_index(result, include_text=True)
    leaf_results = search_index(index, "control testing", top_k=5, leaf_only=True)
    returned_ids = {item.node_id for item in leaf_results}
    assert result["structure"][0]["node_id"] not in returned_ids
